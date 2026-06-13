"""
OCR engine abstraction.

Public API:
    extract_text(file_path, engine="tesseract") -> list[tuple[int | None, str]]

Each tuple is (page_number, text).  page_number is None for single-page
formats (images, plain-text, documents).  Multi-page formats (PDFs,
spreadsheets) yield one tuple per page/sheet.

Supported formats:
  Images (.png .jpg .jpeg .tiff .tif .bmp .gif .webp)
    → Tesseract OCR
  PDF (.pdf)
    → OCRmyPDF + pdftotext; falls back to pdftotext-only for pre-OCR'd PDFs
  Plain text (.txt .md .rst .log .csv .tsv .nfo)
    → read directly; no OCR needed
  Spreadsheets (.xlsx → openpyxl; .xls → xlrd; .ods → odfpy)
    → one result per non-empty sheet (sheet index = page_number)
  Word documents (.docx → python-docx)
    → paragraphs + table cells, page_number=None

All temp files are process-local (/tmp).  Nothing is ever written back
to CORPUS_ROOT or any NFS-mounted path.
"""

import logging
import subprocess
import tempfile
from pathlib import Path

log = logging.getLogger("ocr.engine")

# ─── Extension groups ─────────────────────────────────────────────────────────

IMAGE_EXTENSIONS: frozenset[str] = frozenset({
    ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif", ".webp",
})
PDF_EXTENSION = ".pdf"
PLAIN_TEXT_EXTENSIONS: frozenset[str] = frozenset({
    # Documents / prose
    ".txt", ".md", ".rst", ".log", ".csv", ".tsv", ".nfo",
    # Infrastructure / config
    # ".tf", ".tfvars", ".hcl",
    # ".yaml", ".yml",
    # ".json", ".jsonc",
    # ".toml",
    # ".ini", ".cfg", ".conf",
    # ".xml",
    # ".html", ".htm",
    # # Scripts / code
    # ".sh", ".bash", ".zsh",
    # ".py",
    # ".js", ".ts", ".jsx", ".tsx",
    # ".sql",
    # ".rb", ".go",
    # ".makefile",
})
SPREADSHEET_EXTENSIONS: frozenset[str] = frozenset({
    ".xlsx", ".xls", ".ods",
})
DOCUMENT_EXTENSIONS: frozenset[str] = frozenset({
    ".docx",
})

ALL_EXTENSIONS: frozenset[str] = (
    IMAGE_EXTENSIONS
    | {PDF_EXTENSION}
    | PLAIN_TEXT_EXTENSIONS
    | SPREADSHEET_EXTENSIONS
    | DOCUMENT_EXTENSIONS
)


class OCRError(Exception):
    """Raised when extraction fails for a specific file."""


# ─── Images — Tesseract ───────────────────────────────────────────────────────

def _tesseract(file_path: Path) -> list[tuple[None, str]]:
    try:
        import pytesseract
        from PIL import Image

        with Image.open(file_path) as img:
            text = pytesseract.image_to_string(img)
        return [(None, text)]
    except Exception as e:
        raise OCRError(f"Tesseract failed on {file_path.name}: {e}") from e


# ─── PDFs — OCRmyPDF + pdftotext ─────────────────────────────────────────────

def _pdf_to_pages(pdf_path: Path, timeout: int = 120) -> list[tuple[int, str]]:
    proc = subprocess.run(
        ["pdftotext", "-layout", str(pdf_path), "-"],
        capture_output=True,
        text=True,
        timeout=timeout,
    )
    if proc.returncode != 0:
        raise OCRError(
            f"pdftotext failed (rc={proc.returncode}): {proc.stderr[:500]}")
    pages = []
    for i, chunk in enumerate(proc.stdout.split("\f"), start=1):
        stripped = chunk.strip()
        if stripped:
            pages.append((i, stripped))
    return pages or [(1, proc.stdout.strip())]


def _is_garbled(pages: list[tuple]) -> bool:
    """Return True if the extracted text looks like a garbled custom font encoding.

    Some PDFs use obfuscated glyph-name mappings that cause pdftotext to output
    a single repeated character (commonly 'd') for almost every letter.  When
    one letter accounts for > 70 % of all alphabetic characters the result is
    worthless and we should fall back to rendering + OCR instead.
    """
    letters: list[str] = []
    for _, text in pages:
        letters.extend(c.lower() for c in text if c.isalpha())
    if len(letters) < 100:
        return False  # too short to judge
    top_count = max(letters.count(c) for c in set(letters))
    return top_count / len(letters) > 0.70


def _pdf_via_render_ocr(file_path: Path) -> list[tuple[int, str]]:
    """Render PDF pages to images with pdftoppm and OCR each with Tesseract.

    Used when OCRmyPDF cannot force-OCR the file (e.g. digitally signed PDFs
    with garbled embedded text layers).  Bypasses OCRmyPDF entirely.
    """
    import pytesseract
    from PIL import Image

    with tempfile.TemporaryDirectory() as tmpdir:
        proc = subprocess.run(
            ["pdftoppm", "-png", "-r", "200", str(file_path), f"{tmpdir}/page"],
            capture_output=True,
            timeout=300,
        )
        if proc.returncode != 0:
            raise OCRError(f"pdftoppm failed: {proc.stderr[:300]}")

        pages: list[tuple[int, str]] = []
        for img_path in sorted(Path(tmpdir).glob("page-*.png")):
            try:
                page_num = int(img_path.stem.rsplit("-", 1)[-1])
            except ValueError:
                page_num = len(pages) + 1
            try:
                with Image.open(img_path) as img:
                    text = pytesseract.image_to_string(img, lang="por+deu+eng")
                if text.strip():
                    pages.append((page_num, text.strip()))
            except Exception as e:
                raise OCRError(f"Tesseract failed on page {page_num} of {file_path.name}: {e}") from e

        return pages or [(1, "")]


def _ocrmypdf(file_path: Path) -> list[tuple[int, str]]:
    import ocrmypdf
    import ocrmypdf.exceptions

    with tempfile.TemporaryDirectory() as tmpdir:

        def _run(force: bool) -> list[tuple[int, str]] | None:
            out_pdf = Path(tmpdir) / ("force.pdf" if force else "out.pdf")
            try:
                ocrmypdf.ocr(
                    str(file_path),
                    str(out_pdf),
                    deskew=True,
                    force_ocr=force,
                    skip_text=(not force),
                    progress_bar=False,
                    output_type="pdf",  # Skip PDF/A conversion; avoids color space & Ghostscript crashes
                )
                return _pdf_to_pages(out_pdf)
            except ocrmypdf.exceptions.PriorOcrFoundError:
                # PDF already has a recognised text layer; extract it directly.
                # If that text turns out to be garbled the caller will retry.
                return _pdf_to_pages(file_path)
            except Exception as e:
                # Fall back to direct text extraction from the original PDF if ocrmypdf fails.
                # Crucial for encrypted, signed, or malformed PDFs that already contain character text layers.
                log.warning("ocrmypdf failed on %s: %s. Falling back to direct text extraction.", file_path.name, e)
                try:
                    pages = _pdf_to_pages(file_path)
                    if pages:
                        return pages
                except Exception:
                    pass
                raise OCRError(f"ocrmypdf failed on {file_path.name}: {e}") from e

        pages = _run(force=False)

        if pages and _is_garbled(pages):
            # First try: force_ocr (works for most PDFs with garbled fonts)
            try:
                forced = _run(force=True)
            except OCRError:
                forced = None

            if forced and not _is_garbled(forced):
                pages = forced
            else:
                # force_ocr unavailable or still garbled (e.g. signed PDF).
                # Fall back: render pages as images and OCR with Tesseract directly.
                try:
                    pages = _pdf_via_render_ocr(file_path)
                except OCRError:
                    pass  # keep garbled text rather than failing entirely

        return pages


# ─── Plain text (.txt .csv .md etc.) ─────────────────────────────────────────

def _plain_text(file_path: Path) -> list[tuple[None, str]]:
    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return [(None, text)]
    except Exception as e:
        raise OCRError(f"Failed to read {file_path.name}: {e}") from e


# ─── Spreadsheets ─────────────────────────────────────────────────────────────

def _extract_xlsx(file_path: Path) -> list[tuple[int, str]]:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(str(file_path), read_only=True, data_only=True)
        results = []
        for sheet_idx, ws in enumerate(wb.worksheets, start=1):
            lines = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                line = "\t".join(cells)
                if line.strip():
                    lines.append(line)
            if lines:
                results.append((sheet_idx, "\n".join(lines)))
        wb.close()
        return results
    except Exception as e:
        raise OCRError(f"Failed to read xlsx {file_path.name}: {e}") from e


def _extract_xls(file_path: Path) -> list[tuple[int, str]]:
    try:
        import xlrd

        wb = xlrd.open_workbook(str(file_path))
        results = []
        for sheet_idx, ws in enumerate(wb.sheets(), start=1):
            lines = []
            for row_idx in range(ws.nrows):
                cells = [str(ws.cell_value(row_idx, col))
                         for col in range(ws.ncols)]
                line = "\t".join(cells)
                if line.strip():
                    lines.append(line)
            if lines:
                results.append((sheet_idx, "\n".join(lines)))
        return results
    except Exception as e:
        raise OCRError(f"Failed to read xls {file_path.name}: {e}") from e


def _extract_ods(file_path: Path) -> list[tuple[int, str]]:
    try:
        from odf.opendocument import load
        from odf.table import Table, TableCell, TableRow
        from odf.text import P

        doc = load(str(file_path))
        results = []
        for sheet_idx, sheet in enumerate(
            doc.spreadsheet.getElementsByType(Table), start=1
        ):
            lines = []
            for row in sheet.getElementsByType(TableRow):
                cells = []
                for cell in row.getElementsByType(TableCell):
                    # Collect text content from all paragraph nodes in the cell.
                    cell_text = " ".join(
                        "".join(
                            node.data
                            for node in p.childNodes
                            if hasattr(node, "data")
                        )
                        for p in cell.getElementsByType(P)
                    )
                    cells.append(cell_text)
                line = "\t".join(cells)
                if line.strip():
                    lines.append(line)
            if lines:
                results.append((sheet_idx, "\n".join(lines)))
        return results
    except Exception as e:
        raise OCRError(f"Failed to read ods {file_path.name}: {e}") from e


# ─── Word documents (.docx) ───────────────────────────────────────────────────

def _extract_docx(file_path: Path) -> list[tuple[None, str]]:
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip()
                         for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        return [(None, "\n".join(parts))]
    except Exception as e:
        raise OCRError(f"Failed to read docx {file_path.name}: {e}") from e


# ─── Public entry point ───────────────────────────────────────────────────────

def extract_text(
    file_path: Path,
    engine: str = "tesseract",
) -> list[tuple[int | None, str]]:
    """
    Extract searchable text from file_path.

    Returns a list of (page_number, text) tuples.  page_number is None for
    single-page formats; sheet/page index (1-based) for multi-page formats.

    Raises OCRError on failure or unsupported extension.
    """
    ext = file_path.suffix.lower()

    if ext == PDF_EXTENSION:
        return _ocrmypdf(file_path)

    if ext in IMAGE_EXTENSIONS:
        if engine == "tesseract":
            return _tesseract(file_path)
        raise OCRError(f"Unknown OCR engine: {engine!r}")

    if ext in PLAIN_TEXT_EXTENSIONS:
        return _plain_text(file_path)

    if ext == ".xlsx":
        return _extract_xlsx(file_path)

    if ext == ".xls":
        return _extract_xls(file_path)

    if ext == ".ods":
        return _extract_ods(file_path)

    if ext == ".docx":
        return _extract_docx(file_path)

    raise OCRError(f"Unsupported file type: {ext!r}")
